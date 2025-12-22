import sys
import json
import argparse
import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

class Reference:
    def __init__(self, chrom_list):
        self.chroms = chrom_list
        self.chrom_names = ["chr"+str(i+1) for i in range(22)] + ["chrX", "chrY"]
        self.chr_idx = {self.chroms[c]:c for c in range(len(self.chroms))}
        for c in range(len(self.chrom_names)):
            self.chr_idx[self.chrom_names[c]] = c
        self.chr_len = [None for c in self.chroms]
        self.chr_covg = [None for c in self.chroms]

t2t = Reference(["NC_060925.1", "NC_060926.1", "NC_060927.1", "NC_060928.1", "NC_060929.1", "NC_060930.1", "NC_060931.1", "NC_060932.1", "NC_060933.1", "NC_060934.1", "NC_060935.1", "NC_060936.1", "NC_060937.1", "NC_060938.1", "NC_060939.1", "NC_060940.1", "NC_060941.1", "NC_060942.1", "NC_060943.1", "NC_060944.1", "NC_060945.1", "NC_060946.1", "NC_060947.1", "NC_060948.1"])

def add_text(p, txts, add_newline=True):
    if type(p) == type(Document()):
        p = p.add_paragraph()
    for t in range(len(txts)):
        bold = False
        italic = False
        color = None # default
        if type(txts[t]) == type((0,)):
            txt = txts[t][0]
            bold = txts[t][1]
            if len(txts[t]) >= 3:
                italic = txts[t][2]
            if len(txts[t]) >= 4: # color included as a hex string ex. #FF00AA
                color = txts[t][3]
        else:
            txt = txts[t]
        r = p.add_run(txt)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = RGBColor(
                    int.from_bytes(bytes.fromhex(color[1:3]), byteorder=sys.byteorder),
                    int.from_bytes(bytes.fromhex(color[3:5]), byteorder=sys.byteorder),
                    int.from_bytes(bytes.fromhex(color[5:7]), byteorder=sys.byteorder)
            )
    if add_newline:
        r.add_break()

def bold(cells):
    for c in cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def main(run_dir, outfile, version):
    document = Document()
    sections = document.sections
    margin = 0.5 # inches
    for section in sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    document.add_heading('Adaptive Whole Genome Sequencing Report', 0)

    p = document.add_paragraph()
    add_text(p, [("FOR RESEARCH USE ONLY", True, False, "#FF0000")])

    qf = json.loads(open(f"{run_dir}/fusions.json").read())
    k = json.loads(open(f"{run_dir}/karyotype.json").read())

    now = datetime.datetime.now(datetime.timezone.utc)
    fn = outfile[outfile.rindex('/')+1 if '/' in outfile else 0:]
    p = document.add_paragraph()
    add_text(p, [("Run ID: ", True), f"{fn[:fn.index('.')]}"])
    add_text(p, [("Date/time of report: ", True), f"{now.strftime('%a, %d %b %Y %X UTC')}"])
    add_text(p, [("Analysis version: ", True), f"{version}"])

    document.add_heading("Specimen details", level=2)
    p = document.add_paragraph()
    add_text(p, [("Specimen ID: ", True)])
    add_text(p, [("Date of specimen collection: ", True)])
    add_text(p, [("Specimen source: ", True)], False)

    document.add_heading("Results", level=1)


    # ----------------------- Translocations ----------------------------
    document.add_heading('Gene Fusions', level=2)
    # qf["fusions"] = {... "CBFA2T3-GLIS2": {"supporting_reads": 19, "duplicate_reads": 1, "repetitive_reads": 2} ...}

    blacklist = [("CYP2C19", "CYP2C9"), ("BCR", "IGL"), ("IGK", "IGL"), ("IGK", "KMT2C")]
    filtered_fusions = []
    for fus in qf["fusions"]:
        g0 = fus["gene1"]
        g1 = fus["gene2"]
        if "name" in g0 and "name" in g1 and ((g0["name"], g1["name"]) in blacklist or (g1["name"], g0["name"]) in blacklist):
            continue
        if fus["supporting_reads"] < 3:
            continue
        for b in fus["breakpoints"]:
            if b["n_reads"] > 2 or "DUX4" in g0["name"] or "DUX4" in g1["name"]: # DUX4 is excused from the breakpoint filter
                filtered_fusions.append(fus)
                break
        '''
        "NOL4L-PAX5": {
            "supporting_reads": 15,
            "duplicate_reads": 0, # we don't call dups anymore, so should always be 0
            "repetitive_reads": 7,
            "breakpoints": [
                {
                    "gene0_name": "DUX4L19",
                    "gene0_chr": "NC_060944.1",
                    "gene0_pos": 34173827,
                    "gene0_dir": "<-|",
                    "gene1_name": "DUX4L19",
                    "gene1_chr": "NC_060933.1",
                    "gene1_pos": 36895590,
                    "gene1_dir": "|->",
                    "n_reads": 15
                }
            ]
        }
        '''

    two_sided = [f for f in filtered_fusions if len(f["gene1"]["name"]) > 0 and len(f["gene1"]["name"]) > 0]
    one_sided = [f for f in filtered_fusions if len(f["gene1"]["name"]) == 0 or len(f["gene1"]["name"]) == 0]

    if len(two_sided) > 0:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Gene/region 1'
        hdr_cells[1].text = 'Gene/region 2'
        hdr_cells[2].text = 'Supporting reads'
        hdr_cells[3].text = 'Breakpoint(s)'
        bold(hdr_cells)
        for fus in two_sided:
            g0 = fus["gene1"]
            g1 = fus["gene2"]

            row_cells = table.add_row().cells
            row_cells[0].text = f"{g0['name'] if 'name' in g0 else ''} ({g0['chr']} ~{g0['pos']/100:.2f}Mbp)"
            row_cells[1].text = f"{g1['name'] if 'name' in g1 else ''} ({g1['chr']} ~{g1['pos']/100:.2f}Mbp)"
            row_cells[2].text = f'{fus["supporting_reads"]}'
            if fus["supporting_reads"] >= 3:
                bold([row_cells[2]])
            breakpoints = fus["breakpoints"]
            row_cells[3].text = '\n\r'.join([f'{breakpoints[i]["gene0_chr"]}:{breakpoints[i]["gene0_pos"]}{breakpoints[i]["gene0_dir"]} -- {breakpoints[i]["gene1_chr"]}:{breakpoints[i]["gene1_pos"]}{breakpoints[i]["gene1_dir"]} ({breakpoints[i]["n_reads"]} reads)' for i in range(len(breakpoints))]);
        table.style = "Table Grid"
        document.add_paragraph("")
    else:
        p = document.add_paragraph()
        add_text(p, [(f"No fusions detected", True)])


    document.add_heading('Rearrangements with only one side in the panel', level=3)

    if len(one_sided) > 0:
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Gene/region 1'
        hdr_cells[1].text = 'Gene/region 2'
        hdr_cells[2].text = 'Supporting reads'
        hdr_cells[3].text = 'Breakpoint(s)'
        bold(hdr_cells)
        for fus in one_sided:
            g0 = fus["gene1"]
            g1 = fus["gene2"]

            row_cells = table.add_row().cells
            row_cells[0].text = f"{g0['name'] if 'name' in g0 else ''} ({g0['chr']} ~{g0['pos']/100:.2f}Mbp)"
            row_cells[1].text = f"{g1['name'] if 'name' in g1 else ''} ({g1['chr']} ~{g1['pos']/100:.2f}Mbp)"
            row_cells[2].text = f'{fus["supporting_reads"]}'
            if fus["supporting_reads"] >= 3:
                bold([row_cells[2]])
            breakpoints = fus["breakpoints"]
            row_cells[3].text = '\n\r'.join([f'{breakpoints[i]["gene0_chr"]}:{breakpoints[i]["gene0_pos"]}{breakpoints[i]["gene0_dir"]} -- {breakpoints[i]["gene1_chr"]}:{breakpoints[i]["gene1_pos"]}{breakpoints[i]["gene1_dir"]} ({breakpoints[i]["n_reads"]} reads)' for i in range(len(breakpoints))]);
        table.style = "Table Grid"
        document.add_paragraph("")
    else:
        p = document.add_paragraph()
        add_text(p, [(f"None detected", True)])


    # ----------------------- Karyotype ----------------------------

    #document.add_heading('Digital Karyotype', level=2)
    document.add_heading('Inferred Full Chromosome or Arm Level Gain or Loss', level=2)
    p = document.add_paragraph()
    if 'warning' in k:
        for s in k['warning'].split('\n'):
            add_text(p, [(f"WARNING: ", True, False, "#FF0000"), s])
        add_text(p, ["Digital karyotype: ", (f"{k['karyotype_string']}", True,False,"#FF0000")])
        add_text(p, ["ISCN nomenclature: ", (f"{k['ISCN_karyotype']}", True,False,"#FF0000")])
    else:
        add_text(p, ["Digital karyotype: ", (f"{k['karyotype_string']}", True)])
        add_text(p, ["ISCN nomenclature: ", (f"{k['ISCN_karyotype']}", True)])
    document.add_picture(f"{run_dir}/karyotype.png", width=Inches(7))

    '''
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chromosome'
    #hdr_cells[1].text = 'Accession ID'
    hdr_cells[1].text = 'Copy number'
    hdr_cells[2].text = 'Median depth (reads/Mbp)'
    bold(hdr_cells)
    for chrom in k:
        if chrom in ["ISCN_karyotype", "reads_aligned"]:
            continue
        row_cells = table.add_row().cells
        #row_cells[0].text = t2t.chrom_names[t2t.chroms.index(chrom)]
        row_cells[0].text = chrom
        row_cells[1].text = f'{k[chrom]["n"]}'
        row_cells[2].text = f'{k[chrom]["x"]}'
    table.style = "Table Grid"
    '''


    # ----------------------- CNVs ----------------------------

    document.add_heading('Focal copy-number variation (work in progress)', level=2)
    cnvs = json.loads(open(f"{run_dir}/cnv.json").read())
    #{{"RUNX1": {"focal": 2.699364026432504, "local": 3.908256880733945}, "CDKN2A": {"focal": 0.45934196855635073, "local": 0.9031600407747197}, "ERG": {"deletions": [{"start": 36773422, "end": 36822979, "reads": 12}]}}}

    document.add_heading('RUNX1', level=3)
    p = document.add_paragraph()
    if "RUNX1" in cnvs:
        add_text(p, [f"Local read-depth based copy number estimate: ", (f"{int(round(cnvs['RUNX1']['local']))}x ({cnvs['RUNX1']['local']:.1f}x)", True)])
    else:
        add_text(p, "Not tested")

    document.add_heading('CDKN2A', level=3)
    p = document.add_paragraph()
    if "CDKN2A" in cnvs:
        add_text(p, [f"Focal sequencing-depth based copy number estimate: ", (f"{int(round(cnvs['CDKN2A']['focal']))}x ({cnvs['CDKN2A']['focal']:.1f}x)", True)])
    else:
        add_text(p, "Not tested")

    document.add_heading('CDKN2B', level=3)
    p = document.add_paragraph()
    if "CDKN2B" in cnvs:
        add_text(p, [f"Focal sequencing-depth based copy number estimate: ", (f"{int(round(cnvs['CDKN2B']['focal']))}x ({cnvs['CDKN2B']['focal']:.1f}x)", True)])
    else:
        add_text(p, "Not tested")

    document.add_heading('ERG intragenic deletions', level=3)
    p = document.add_paragraph()
    if "ERG" in cnvs and len(cnvs["ERG"]["deletions"]) > 0:
        for d in cnvs["ERG"]["deletions"]:
            add_text(p, [f"{d['chrom']}: {d['start']} - {d['end']} ({d['end']-d['start']} nt): ", (f"{d['reads']} supporting reads", True)])
    else:
        add_text(p, "None detected")

    document.add_heading('IKZF1 intragenic deletions', level=3)
    p = document.add_paragraph()
    if "IKZF1" in cnvs and len(cnvs["IKZF1"]["deletions"]) > 0:
        for d in cnvs["IKZF1"]["deletions"]:
            add_text(p, [f"{d['chrom']}: {d['start']} - {d['end']} ({d['end']-d['start']} nt): ", (f"{d['reads']} supporting reads", True)])
    else:
        add_text(p, "None detected")

        document.add_heading('IKZF1', level=3)
        p = document.add_paragraph()
        if "IKZF1" in cnvs:
            add_text(p, [f"Focal sequencing-depth based copy number estimate: ", (f"{int(round(cnvs['IKZF1']['focal']))}x  ({cnvs['IKZF1']['focal']:.1f}x)", True)])
        else:
            add_text(p, "Not tested")

    document.add_heading('PAX5 intragenic deletions', level=3)
    p = document.add_paragraph()
    if "PAX5" in cnvs and len(cnvs["PAX5"]["deletions"]) > 0:
        for d in cnvs["PAX5"]["deletions"]:
            add_text(p, [f"{d['chrom']}: {d['start']} - {d['end']} ({d['end']-d['start']} nt): ", (f"{d['reads']} supporting reads", True)])
    else:
        add_text(p, "None detected")

    document.add_heading('PAX5 intragenic tandem duplications', level=3)
    p = document.add_paragraph()
    if "PAX5" in cnvs and len(cnvs["PAX5"]["duplications"]) > 0:
        for d in cnvs["PAX5"]["duplications"]:
            add_text(p, [f"{d['chrom']}: {d['start']} - {d['end']} ({d['end']-d['start']} nt): ", (f"{d['reads']} supporting reads", True)])
    else:
        add_text(p, "None detected")

    document.add_heading('KMT2A intragenic tandem duplications', level=3)
    p = document.add_paragraph()
    if "KMT2A" in cnvs and len(cnvs["KMT2A"]["duplications"]) > 0:
        for d in cnvs["KMT2A"]["duplications"]:
            add_text(p, [f"{d['chrom']}: {d['start']} - {d['end']} ({d['end']-d['start']} nt): ", (f"{d['reads']} supporting reads", True)])
    else:
        add_text(p, "None detected - note that KMT2A partial tandem duplications that are not entirely within the gene will be reported as fusions of KMT2A with a proximal region of the genome")


    # ----------------------- FLT3-ITD ----------------------------

    document.add_heading('Internal/partial tandem duplications', level=2)
    itds = json.loads(open(f"{run_dir}/itd.json").read())

    document.add_heading('FLT3 internal tandem duplication (ITD)', level=3)
    p = document.add_paragraph()
    if len(itds["FLT3"]) == 0:
        add_text(p, "None detected")
    else:
        for ins in itds["FLT3"]:
            add_text(p, f"{ins['length']} nt insertion at position {ins['position']}: {ins['merged']} reads (of {ins['coverage']}); {ins['merged']/(ins['coverage']-ins['merged']):.2f} AR")


    # ----------------------- UBTF-ITD ----------------------------

    document.add_heading('UBTF internal tandem duplication (ITD)', level=3)
    p = document.add_paragraph()
    if len(itds["UBTF"]) == 0:
        add_text(p, "None detected")
    else:
        for ins in itds["UBTF"]:
            add_text(p, f"{ins['length']} nt insertion at position {ins['position']}: {ins['merged']} reads (of {ins['coverage']}); {ins['merged']/(ins['coverage']-ins['merged']):.2f} AR")


    # ----------------------- KMT2A-ITD ----------------------------

    #document.add_heading('KMT2A partial tandem duplication (PTD)', level=3)
    #document.add_paragraph("TBD")


    # ----------------------- SNVs / genotypes ----------------------------

    gens = json.loads(open(f"{run_dir}/genotypes.json").read())
    # coverage, mutations, phase, genotype
    document.add_heading('SNV Genotypes', level=2)
    document.add_heading('Disease variants', level=3)
    for g in ["PAX5", "IKZF1"]:
        p = document.add_paragraph()
        add_text(p, [(f"{g}", True)])
        add_text(p, [f"Sequencing depth: {gens[g]['coverage']:.2f}x"])
        add_text(p, [f"Mutations:"])
        for m in gens[g]['mutations']:
            add_text(p, [f"  {m}: {gens[g]['mutations'][m]} AF"])
        if len(gens[g]['mutations']) == 0:
            add_text(p, ["(none detected)"])
        add_text(p, [f"Genotype: ", (f"{gens[g]['genotype']}", True)], False)
    document.add_heading('Pharmacogenomic variants', level=3)
    for g in ["TPMT", "NUDT15"]:
        p = document.add_paragraph()
        add_text(p, [(f"{g}", True)])
        add_text(p, [f"Sequencing depth: {gens[g]['coverage']:.2f}x"])
        add_text(p, [f"Mutations:"])
        for m in gens[g]['mutations']:
            add_text(p, [f"  {m}: {gens[g]['mutations'][m]} AF"])
        if len(gens[g]['mutations']) == 0:
            add_text(p, ["(none detected)"])
        add_text(p, [f"Genotype: ", (f"{gens[g]['genotype']}", True)], False)


    # ------------------------ QC -----------------------------

    document.add_heading("Quality control metrics:", level=2)

    document.add_paragraph("Blast percentage reported by hematopathology: ", style="List Bullet")
    document.add_paragraph(f"Genome-wide reads aligned: {int(float(k['reads_aligned'])):,}", style="List Bullet")
    document.add_paragraph(f"Mean on-target read length: {int(qf['qc']['nt_on_target'] / qf['qc']['reads_on_target'])} nt", style="List Bullet")
    document.add_paragraph(f"Mean coverage over target regions: {qf['qc']['nt_on_target']/qf['qc']['target_regions_nt']:.2f}x", style="List Bullet")
    #document.add_paragraph(f"Relative enrichment: {qf['qc']['nt_on_target']/qf['qc']['target_regions_nt']/(qf['qc']['nt_aligned']/3.1e9):.2f}x", style="List Bullet")

    if "meta" in qf and "run_start_time" in qf["meta"]:
        document.add_heading("Sequencing details", level=2)
        p = document.add_paragraph()
        add_text(p, [("Date/time of sequencing start: ", True), qf["meta"]["run_start_time"]])
        add_text(p, [("Sequencing run ID: ", True), qf["meta"]["run_id"]])
        add_text(p, [("Basecalling model: ", True), qf["meta"]["basecall_model"]])
        add_text(p, [("Library ID: ", True), qf["meta"]["library_id"]])
        add_text(p, [("Sequencer ID: ", True), qf["meta"]["sequencer_id"]])
        add_text(p, [("Flow cell ID: ", True), qf["meta"]["flow_cell_id"]], False)


    # ------------------------- horizontal rule, then Methodology -------------------------
    p = document.add_paragraph()
    insertHR(p)
    document.add_heading('Methodology', level=2)

    p = document.add_paragraph()
    add_text(p, [("Digital Karyotype and Copy Number Variants: ", True), f"Relative copy number across the genome at the chromosome level ('digital karyotype') is inferred based on relative sequencing depth by assessing genome-wide and chromosome-level depth of coverage. To avoid the potentially confounding effect of adaptive sampling on relative sequencing depth assessment, coarse-scale depth is constructed as a function of reads per million base pairs (Mbp), where each read contributes a count of one to the bin in which the center of the read aligns. A pairwise Kolmogorov-Smirnov test is applied with a threshold of p < 1e-9 and a minimum median divergence of 20% to determine which chromosomes exhibit significantly divergent copy number."], False)

    p = document.add_paragraph()
    add_text(p, [("Internal Tandem Duplications", True)])
    add_text(p, [("FLT3:", True), f"Insertions are called among reads aligning to FLT3 exons 14 or 15, then called insertions are clustered together if their respective lengths are within 10% of one another and their position in the reference is within 110% of the length of the insertion from one another. Allelic ratio (AR) is calculated as the ratio of the number of reads supporting the insertion cluster divided by the number of reads aligning to the insertion site that are not in the insertion cluster."])
    add_text(p, [("UBTF:", True), f"Insertions are called among reads aligning to UBTF exon 13, then called insertions are clustered together if their respective lengths are within 10% of one another and their position in the reference is within 110% of the length of the insertion from one another. Allelic ratio (AR) is calculated as the ratio of the number of reads supporting the insertion cluster divided by the number of reads aligning to the insertion site that are not in the insertion cluster."], False)

    p = document.add_paragraph()
    add_text(p, [("Single Nucleotide Variants", True)])
    add_text(p, ["Pharmacogenomics: TPMT and NUDT15"])
    add_text(p, ["B-ALL subtype defining:"])
    add_text(p, ["PAX5 c.239C>G, p.Pro80Arg"])
    add_text(p, ["IKZF1 c.475A>T, p.Asm159Tyr"])

    p = document.add_paragraph()
    add_text(p, [("Gene / Regions Enrichment set:", True)])
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    row = 0
    col = -1
    for gene in qf["qc"]["gene_coverage"]:
        col += 1
        if col >= 6:
            col = 0
            row += 1
            table.add_row().cells
        table.rows[row].cells[col].text = gene

    document.save(outfile)

if __name__ == "__main__":
    parser = argparse.ArgumentParser("Make WGS report (docx) from analysis logs/directory")
    parser.add_argument("dir", help="Analysis output directory")
    parser.add_argument("out", help="Output log file (.docx)")
    parser.add_argument("version", help="Analysis version (string)")
    args = parser.parse_args()
    main(args.dir, args.out, args.version)
